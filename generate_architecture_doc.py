"""Generate updated system architecture document (docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')


# ============ TITLE ============
title = doc.add_heading('Multimodal Biometric Authentication System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Updated System Architecture\nKeystroke Dynamics + Speech Biometrics (2FA Alternative)')
r.italic = True
r.font.size = Pt(12)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Document Date: April 29, 2026\nProof of Concept — Desktop Web Only').italic = True

doc.add_page_break()

# ============ 1. OVERVIEW ============
add_heading('1. System Overview', level=1)
add_para(
    "This system is a multimodal biometric authentication proof of concept that combines "
    "keystroke dynamics and speech biometrics as an alternative to traditional 2FA. "
    "It targets desktop web browsers and uses fixed lowercase 4-word passphrases (no mobile, "
    "shift-key, or IME support). Authentication is performed sequentially across four steps: "
    "password, keystroke, voice, and a fusion stage that resolves uncertain keystroke decisions."
)

add_heading('1.1 Core Goals', level=2)
add_bullet('Replace SMS/OTP-based 2FA with behavioral and physiological biometrics.')
add_bullet('Maintain low FAR/FRR even with limited enrollment data via progressive thresholds.')
add_bullet('Continuously adapt to genuine user drift through adaptive sample retention and retraining.')
add_bullet('Resist impostor pollution via pending-sample caching and CMU-based impostor patching.')

# ============ 2. AUTHENTICATION FLOW ============
add_heading('2. Authentication Flow', level=1)
add_para(
    "The login process is a four-step sequential pipeline. The voice stage is skipped entirely "
    "when keystroke confidence is high enough (≥ 0.80), which keeps the user experience fast for "
    "well-enrolled users while preserving an additional verification factor for borderline cases."
)

add_heading('Step 1 — Password', level=2)
add_para('Endpoint: POST /api/auth/password')
add_bullet('Bcrypt verification against the stored password hash.')
add_bullet('No biometric data is touched if the password is wrong.')

add_heading('Step 2 — Keystroke Dynamics', level=2)
add_para('Endpoint: POST /api/auth/keystroke')
add_bullet('Frontend captures dwell, flight, digraph/trigraph (DD/DU/UD/UU variants), per-key dwell map, and quality score.')
add_bullet('Backend scores via Random Forest / GBM (mature users) or Profile Matcher (≤ 10 samples).')
add_bullet('Score is fused with a Mahalanobis distance term using utils/fusion.py weights.')
add_bullet('Confidence ≥ 0.80 → access granted, voice skipped. 0.55–0.79 → uncertain, proceed to fusion. < 0.55 → reject.')

add_heading('Step 3 — Voice (Speaker + Phrase)', level=2)
add_para('Endpoint: POST /api/auth/voice')
add_bullet('ECAPA-TDNN (SpeechBrain pretrained, 192-dim) computes cosine similarity vs. enrolled embedding.')
add_bullet('Whisper transcribes audio for phrase verification (must match the assigned 4-word phrase).')
add_bullet('RNNoise denoising is applied client-side before features are sent.')

add_heading('Step 4 — Multimodal Fusion', level=2)
add_para('Endpoint: POST /api/auth/fuse')
add_bullet('Triggered only when keystroke confidence falls in the uncertain band (0.55–0.79).')
add_bullet('Weighted combination of keystroke and voice scores; weights centralized in utils/fusion.py.')
add_bullet('Successful fusion persists pending keystroke samples held in memory; failed fusion discards them.')

# ============ 3. BACKEND ============
add_heading('3. Backend Architecture (FastAPI)', level=1)

add_heading('3.1 Module Layout', level=2)
add_bullet('main.py — App setup, CORS, static file serving (no-cache headers), router mounts. Sets HF_HUB_DISABLE_SYMLINKS=1 to avoid Windows WinError 1314.')
add_bullet('routers/enroll.py — User creation (bcrypt password, Fernet-encrypted phrase), keystroke/voice/security-question enrollment, audio feature extraction (/extract-mfcc).')
add_bullet('routers/auth.py — Password / keystroke / voice / security verification, multimodal fusion, progressive enrollment thresholds, adaptive keystroke sample saving.')
add_bullet('utils/fusion.py — Single source of truth for intra-modal (RF + Mahalanobis) and inter-modal (keystroke + voice) score weights.')
add_bullet('utils/crypto.py — Fernet encryption/decryption for users.phrase and security_questions.question.')
add_bullet('database/db.py — SQLAlchemy engine with Supabase-tuned pool settings (pool_pre_ping, keepalives).')
add_bullet('database/models.py — User, KeystrokeTemplate (60+ feature columns + JSON digraph variant maps), VoiceTemplate, SecurityQuestion, AuthLog.')
add_bullet('schemas.py — VoiceFeatures base shared by enrollment and auth (single source of truth for voice fields).')

add_heading('3.2 Key Design Patterns', level=2)
add_bullet('Progressive enrollment: keystroke threshold begins soft (0.45) for new users, ramps to 0.55 after 3 successful logins, hardens to the stored threshold after 7. Maturity = count of granted keystroke/fusion auth logs.')
add_bullet('Adaptive learning: successful keystroke samples are saved back to the DB and trigger periodic retraining (every 2 / 5 / 10 logins depending on sample count, capped at 50 samples with FIFO eviction).')
add_bullet('Pending sample cache: keystroke samples from sessions that fail keystroke but pass via voice fusion are held in _pending_ks_samples and only persisted if fusion grants access — prevents impostor or noisy samples from polluting training data.')
add_bullet('Auto-training: enrollment triggers background-thread training after 5 keystroke or 3 voice samples have been collected.')
add_bullet('Phrase-aware features: only digraphs present in the assigned passphrase are kept as features; inactive hardcoded digraphs are stripped before training.')

# ============ 4. ML LAYER ============
add_heading('4. Machine Learning Layer', level=1)

add_heading('4.1 Keystroke Pipeline', level=2)
add_bullet('train_keystroke_rf.py — Phrase-aware training pipeline: strips inactive digraphs, generates 3-tier synthetic impostors (near / shifted / random), patches with CMU impostor profiles, performs FAR-weighted threshold search.')
add_bullet('keystroke_profile_matcher.py — TypingDNA-style direct comparison with no classifier: speed normalization, tiered Z-tolerance per feature group, weighted group scoring. Used for users with ≤ 10 samples.')
add_bullet('Upgrade rule: at 11+ samples the user transitions from Profile Matcher to Random Forest / GBM scoring.')
add_bullet('load_cmu_impostors.py — Builds cmu_impostor_profiles.pkl from the CMU keystroke dataset for impostor augmentation.')

add_heading('4.2 Voice Pipeline', level=2)
add_bullet('voice_ecapa.py — ECAPA-TDNN speaker verification (SpeechBrain pretrained), 192-dim embeddings, cosine similarity scoring.')
add_bullet('train_voice_cnn.py — Auxiliary CNN voice model trainer.')
add_bullet('Phrase verification via Whisper transcription against the user’s assigned 4-word passphrase.')

add_heading('4.3 Model Storage', level=2)
add_bullet('Models saved as .pkl files in ml/models/ with sanitized filenames (@ → _at_, . → _).')
add_bullet('Per-user keystroke and voice models; retraining writes back to the same path atomically.')

# ============ 5. FRONTEND ============
add_heading('5. Frontend Architecture', level=1)
add_bullet('js/keystroke.js — Captures dwell times, flight times, digraph/trigraph maps (DD/DU/UD/UU variants), per-key dwell map, and a typing quality score.')
add_bullet('js/speech.js — Microphone capture with RNNoise denoising; sends base64-encoded audio to /extract-mfcc.')
add_bullet('js/api.js — Auto-detects API base URL from window.location; shared voice payload builder for enroll and auth.')
add_bullet('js/enroll.js — Enrollment flow orchestration (password, phrase, keystroke samples, voice samples, security questions).')
add_bullet('js/auth_flow.js — Login flow orchestration (sequential password → keystroke → voice → fusion).')
add_bullet('Pages: login.html, enroll.html, dashboard.html, forgot-password.html, reset.html.')

# ============ 6. DATABASE ============
add_heading('6. Database (PostgreSQL on Supabase)', level=1)
add_bullet('users.phrase column stores Fernet-encrypted ciphertext — callers MUST decrypt() before using plaintext.')
add_bullet('security_questions.question is also Fernet-encrypted.')
add_bullet('Credentials provided via .env (DB_USER, DB_PASSWORD, DB_HOST, DB_PORT, DB_NAME, ENCRYPTION_KEY).')
add_bullet('Schema migrations are applied manually via SQL — see comments in database/models.py for ALTER TABLE statements. No migration framework.')

add_heading('6.1 Core Tables', level=2)
add_bullet('User — credentials, encrypted phrase, enrollment maturity counters.')
add_bullet('KeystrokeTemplate — 60+ feature columns plus JSON maps for the four digraph timing variants.')
add_bullet('VoiceTemplate — ECAPA embedding(s), enrollment metadata.')
add_bullet('SecurityQuestion — encrypted question + answer hash for password recovery.')
add_bullet('AuthLog — granular event log used to compute maturity and drive progressive thresholds.')

# ============ 7. RUNTIME ENVIRONMENT ============
add_heading('7. Runtime Environment', level=1)
add_bullet('Two virtual environments: venv (primary) and venv310 (Python 3.10 for torch / SpeechBrain compatibility).')
add_bullet('start.bat offers (1) stable/demo mode and (2) dev mode with auto-reload.')
add_bullet('Manual launch: uvicorn main:app --host 127.0.0.1 --port 8000 [--reload].')
add_bullet('Static frontend served by FastAPI with no-cache headers to avoid stale JS during iteration.')

# ============ 8. SECURITY POSTURE ============
add_heading('8. Security Posture', level=1)
add_bullet('Passwords hashed with bcrypt; phrases and security questions encrypted at rest with Fernet.')
add_bullet('Pending-sample caching prevents impostor or noisy keystroke samples from being persisted unless a second factor (voice fusion) confirms identity.')
add_bullet('FAR-weighted threshold search during training biases toward lower False Accept Rate.')
add_bullet('CMU impostor patching exposes models to realistic attacker keystroke distributions during training.')
add_bullet('Audit trail: every authentication attempt and decision recorded in AuthLog.')

# ============ 9. END-TO-END SEQUENCE ============
add_heading('9. End-to-End Login Sequence', level=1)
add_para('1. User submits email + password → bcrypt verified.')
add_para('2. Frontend records keystroke timing of the assigned passphrase → POST /api/auth/keystroke.')
add_para('3. Backend scores via Profile Matcher (early enrollment) or RF/GBM (mature) + Mahalanobis fusion.')
add_para('4. If confidence ≥ 0.80 → grant access, log event, optionally save sample for adaptive learning, end.')
add_para('5. If confidence < 0.55 → reject and log failure.')
add_para('6. If 0.55 ≤ confidence < 0.80 → cache keystroke sample as pending, prompt voice capture.')
add_para('7. Voice ECAPA cosine + Whisper phrase check → POST /api/auth/voice → POST /api/auth/fuse.')
add_para('8. Weighted keystroke + voice fusion ≥ threshold → grant access, persist pending sample, retrain if eligible. Otherwise reject and discard pending sample.')

# ============ 10. FUTURE WORK ============
add_heading('10. Notes & Constraints', level=1)
add_bullet('PoC scope: desktop web only; no mobile, no shift-key handling, no IME.')
add_bullet('Fixed lowercase 4-word passphrase format keeps the keystroke feature space tractable.')
add_bullet('Weights and thresholds are centralized in utils/fusion.py to keep tuning auditable.')
add_bullet('Two voice models may be in active training at any time; retraining is throttled by login count tiers.')

# ============ SAVE ============
output_path = r'c:\Users\John Lorenz\Desktop\biometric_auth\System_Architecture_Updated.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
